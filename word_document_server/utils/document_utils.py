"""
Document utility functions for Word Document Server.
"""
import json
from typing import Dict, List, Any
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Skip TOC paragraphs in tables
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count


def get_document_xml(doc_path: str) -> str:
    """Get the full XML of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        return doc.element.body.xml
    except Exception as e:
        return f"Failed to get document XML: {str(e)}"

def get_paragraph_xml_by_text(doc_path: str, search_text: str) -> str:
    """Get the XML of a paragraph containing specific text."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        for para in doc.paragraphs:
            if search_text in para.text:
                return para._p.xml
        return f"Paragraph containing '{search_text}' not found."
    except Exception as e:
        return f"Failed to get paragraph XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def insert_numbered_list_near_text(doc_path: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None) -> str:
    """
    Insert a numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Robust style selection for numbered list
        style_name = None
        for candidate in ['List Number', 'List Paragraph', 'Normal']:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default
        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Numbered list inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Numbered list inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text (like delete_paragraph finds by index)
    header_para = None
    header_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == header_text.strip().lower():
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first
    header_para = None
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().lower()
        is_toc = is_toc_paragraph(para)
        if para_text == header_text.strip().lower() and not is_toc:
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor
    for i, el in enumerate(elements):
        if el.tag == CT_P.tag:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif p_text == start_anchor_text.strip():
                start_idx = i
                break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif p_text == end_anchor_text.strip():
                    end_idx = i
                    break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if para.text.strip() == start_anchor_text.strip():
            anchor_idx = i
            break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."


def clone_run_with_text(source_run, text):
    """
    Clone a run with all its formatting but replace the text content.
    
    Args:
        source_run: The original run to clone formatting from
        text: New text content for the cloned run
        
    Returns:
        New run element with same formatting as source but different text
    """
    import copy
    from docx.oxml.ns import qn
    
    # Deep copy the run element to preserve all formatting
    cloned_run_element = copy.deepcopy(source_run._r)
    
    # Update the text content in all text nodes
    for t_elem in cloned_run_element.iter(qn('w:t')):
        t_elem.text = text
        break  # Only set text in the first text element, clear others
    
    # Clear text from remaining text elements
    text_elements = list(cloned_run_element.iter(qn('w:t')))
    for t_elem in text_elements[1:]:
        t_elem.text = ""
    
    return cloned_run_element


def split_run_at(run, split_index):
    """
    Split a run at the specified character index.
    
    Args:
        run: The run to split
        split_index: Character index where to split (0-based)
        
    Returns:
        Tuple of (left_text, right_text) - the text portions after split
    """
    original_text = run.text
    if split_index <= 0:
        return "", original_text
    if split_index >= len(original_text):
        return original_text, ""
    
    left_text = original_text[:split_index]
    right_text = original_text[split_index:]
    
    # Update the original run with left text
    run.text = left_text
    
    # Create a new run with right text and same formatting
    right_run_element = clone_run_with_text(run, right_text)
    
    # Insert the right run immediately after the original
    run._r.addnext(right_run_element)
    
    return left_text, right_text


def normalize_text_for_search(text):
    """
    Normalize text by removing invisible separators and converting special characters.
    
    Args:
        text: Text to normalize
        
    Returns:
        Normalized text for search matching
    """
    if not text:
        return text
        
    # Remove zero-width characters and soft hyphens
    text = text.replace('\u200b', '')  # Zero-width space
    text = text.replace('\u200c', '')  # Zero-width non-joiner
    text = text.replace('\u200d', '')  # Zero-width joiner
    text = text.replace('\u00ad', '')  # Soft hyphen
    text = text.replace('\ufeff', '')  # Zero-width no-break space (BOM)
    
    # Convert NBSP to regular space
    text = text.replace('\u00a0', ' ')
    
    return text


def build_xml_run_map(para):
    """
    Build run mapping using XML traversal to catch runs inside wrappers.
    This handles cases where runs are nested in content controls, hyperlinks, etc.
    
    Args:
        para: Paragraph object
        
    Returns:
        Tuple of (full_text, xml_run_map) where xml_run_map contains 
        (start_pos, end_pos, run_element) for each run found via XML
    """
    from docx.oxml.ns import qn
    
    full_text = ""
    xml_run_map = []
    
    # Find all w:r elements under this paragraph using element iteration
    def find_all_runs(element):
        """Recursively find all run elements."""
        runs = []
        if element.tag == qn('w:r'):
            runs.append(element)
        for child in element:
            runs.extend(find_all_runs(child))
        return runs
    
    run_elements = find_all_runs(para._p)
    
    for run_element in run_elements:
        start_pos = len(full_text)
        run_text = ""
        
        # Extract text from all w:t elements in this run
        for t_elem in run_element.iter(qn('w:t')):
            if t_elem.text:
                run_text += t_elem.text
        
        # Handle w:tab elements as spaces  
        tab_count = len(list(run_element.iter(qn('w:tab'))))
        run_text += ' ' * tab_count
        
        if run_text:  # Only add if run has text
            normalized_text = normalize_text_for_search(run_text)
            full_text += normalized_text
            end_pos = len(full_text)
            xml_run_map.append((start_pos, end_pos, run_element))
    
    return full_text, xml_run_map


def find_and_replace_text_preserve_formatting(doc, old_text, new_text):
    """
    Find and replace text throughout the document while preserving original formatting.
    Handles mixed formatting within the matched text (e.g., some parts bold, some italic).
    Now includes XML-based fallback for runs nested in wrappers (content controls, hyperlinks).
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0

    # Lightweight debug logging (enable by setting env MCP_WORD_DEBUG_FIND_REPLACE=1)
    import os
    DEBUG_FIND_REPLACE = os.getenv('MCP_WORD_DEBUG_FIND_REPLACE', '0') == '1'
    def _dbg(msg):
        if DEBUG_FIND_REPLACE:
            print(f"[find_replace] {msg}")
    
    def process_paragraph(para):
        """Process a single paragraph for search and replace with formatting preservation."""
        nonlocal count
        
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            return
            
        # First try: Build flattened text using standard paragraph.runs API
        full_text = ""
        run_map = []  # List of (start_pos, end_pos, run_index)
        
        for run_idx, run in enumerate(para.runs):
            start_pos = len(full_text)
            normalized_run_text = normalize_text_for_search(run.text)
            full_text += normalized_run_text
            end_pos = len(full_text)
            if end_pos > start_pos:  # Only add if run has text
                run_map.append((start_pos, end_pos, run_idx))
        
        # Normalize search text for consistent matching
        normalized_old_text = normalize_text_for_search(old_text)
        
        # Check if we can find the text using standard API (even if full_text == "")
        if normalized_old_text in full_text:
            _dbg("Standard path: match found in paragraph.runs")
            process_paragraph_standard(para, full_text, run_map, normalized_old_text, new_text)
        else:
            _dbg("Standard path: no match or no direct runs; trying XML fallback")
            # Fallback: Try XML-based approach for runs inside wrappers
            xml_full_text, xml_run_map = build_xml_run_map(para)
            if normalized_old_text in xml_full_text:
                _dbg("XML fallback: match found in nested runs")
                process_paragraph_xml(para, xml_full_text, xml_run_map, normalized_old_text, new_text)
            else:
                _dbg("XML fallback: no match in this paragraph")
    
    def process_paragraph_standard(para, full_text, run_map, old_text_norm, new_text):
        """Process paragraph using standard paragraph.runs API."""
        nonlocal count
        
        if not full_text or old_text_norm not in full_text:
            return
            
        # Find all occurrences (process right to left to avoid index shifts)
        matches = []
        start_search = 0
        while True:
            match_pos = full_text.find(old_text_norm, start_search)
            if match_pos == -1:
                break
            matches.append((match_pos, match_pos + len(old_text_norm)))
            start_search = match_pos + 1
        
        _dbg(f"Standard path: {len(matches)} match(es) in paragraph")
        
        # Process matches from right to left to avoid index corruption
        for match_start, match_end in reversed(matches):
            # Find which runs overlap with this match
            affected_runs = []
            for start_pos, end_pos, run_idx in run_map:
                if start_pos < match_end and end_pos > match_start:
                    # Calculate overlap within this run
                    overlap_start = max(0, match_start - start_pos)
                    overlap_end = min(end_pos - start_pos, match_end - start_pos)
                    affected_runs.append((run_idx, overlap_start, overlap_end, start_pos))
            
            if not affected_runs:
                continue
                
            # Process the replacement
            runs = para.runs
            first_run_idx, first_overlap_start, first_overlap_end, first_run_start = affected_runs[0]
            last_run_idx, last_overlap_start, last_overlap_end, last_run_start = affected_runs[-1]
            
            # Split boundary runs if needed
            first_run = runs[first_run_idx]
            if first_overlap_start > 0:
                # Split the first run, keeping the left part
                split_run_at(first_run, first_overlap_start)
                # The right part becomes the new "first run" for our match
                first_run = runs[first_run_idx + 1] if first_run_idx + 1 < len(runs) else first_run
                # Update indices for all subsequent runs
                for i in range(len(affected_runs)):
                    if affected_runs[i][0] > first_run_idx:
                        affected_runs[i] = (affected_runs[i][0] + 1, affected_runs[i][1], affected_runs[i][2], affected_runs[i][3])
                first_run_idx += 1
                last_run_idx += 1
            
            # Re-get runs after potential split
            runs = para.runs
            last_run = runs[last_run_idx] if last_run_idx < len(runs) else runs[-1]
            
            # If last run extends beyond match, split it
            last_run_text = last_run.text
            if last_overlap_end < len(last_run_text):
                split_run_at(last_run, last_overlap_end)
            
            # Re-get runs after potential split
            runs = para.runs
            
            # Collect the runs that need to be replaced (the matched portion)
            runs_to_replace = []
            for run_idx in range(first_run_idx, min(last_run_idx + 1, len(runs))):
                runs_to_replace.append(runs[run_idx])
            
            if not runs_to_replace:
                continue
                
            # Create replacement runs by cloning original formatting
            replacement_runs = []
            remaining_text = new_text
            
            for i, original_run in enumerate(runs_to_replace):
                if not remaining_text:
                    break
                    
                # Determine how much text to put in this run segment
                if i == len(runs_to_replace) - 1:
                    # Last segment gets all remaining text
                    segment_text = remaining_text
                    remaining_text = ""
                else:
                    # Distribute text proportionally, but at least 1 char if text remains
                    original_segment_len = len(original_run.text)
                    if original_segment_len > 0:
                        chars_to_take = min(original_segment_len, len(remaining_text))
                    else:
                        chars_to_take = 1 if remaining_text else 0
                    
                    segment_text = remaining_text[:chars_to_take]
                    remaining_text = remaining_text[chars_to_take:]
                
                if segment_text:  # Only create run if there's text
                    cloned_run_element = clone_run_with_text(original_run, segment_text)
                    replacement_runs.append(cloned_run_element)
            
            # If there's still remaining text, create additional runs using last run's formatting
            while remaining_text and runs_to_replace:
                last_original = runs_to_replace[-1]
                cloned_run_element = clone_run_with_text(last_original, remaining_text)
                replacement_runs.append(cloned_run_element)
                remaining_text = ""
            
            # Replace the matched runs with replacement runs
            if replacement_runs and runs_to_replace:
                # Insert replacement runs before the first run to replace
                insert_point = runs_to_replace[0]._r
                for replacement_run in replacement_runs:
                    insert_point.addprevious(replacement_run)
                
                # Remove the original matched runs
                for run_to_remove in runs_to_replace:
                    run_to_remove._r.getparent().remove(run_to_remove._r)
            
            count += 1
            _dbg("Standard path: one replacement applied")
    
    def process_paragraph_xml(para, xml_full_text, xml_run_map, old_text_norm, new_text):
        """Process paragraph using XML-based run discovery for nested runs."""
        nonlocal count
        
        if not xml_full_text or old_text_norm not in xml_full_text:
            return
            
        # Find all occurrences (process right to left to avoid index shifts)
        matches = []
        start_search = 0
        while True:
            match_pos = xml_full_text.find(old_text_norm, start_search)
            if match_pos == -1:
                break
            matches.append((match_pos, match_pos + len(old_text_norm)))
            start_search = match_pos + 1
        
        _dbg(f"XML path: {len(matches)} match(es) in paragraph")
        
        # Process matches from right to left to avoid index corruption
        for match_start, match_end in reversed(matches):
            # Find which XML runs overlap with this match
            affected_xml_runs = []
            for start_pos, end_pos, run_element in xml_run_map:
                if start_pos < match_end and end_pos > match_start:
                    # Calculate overlap within this run
                    overlap_start = max(0, match_start - start_pos)
                    overlap_end = min(end_pos - start_pos, match_end - start_pos)
                    affected_xml_runs.append((run_element, overlap_start, overlap_end, start_pos))
            
            if not affected_xml_runs:
                continue
                
            # For XML-based approach, we'll do a simpler replacement:
            # Clone the formatting from affected runs and create new runs
            replacement_runs = []
            remaining_text = new_text
            
            for i, (run_element, overlap_start, overlap_end, run_start_pos) in enumerate(affected_xml_runs):
                if not remaining_text:
                    break
                    
                # Determine how much text to put in this run segment
                if i == len(affected_xml_runs) - 1:
                    # Last segment gets all remaining text
                    segment_text = remaining_text
                    remaining_text = ""
                else:
                    # Distribute text proportionally
                    original_segment_len = overlap_end - overlap_start
                    if original_segment_len > 0:
                        chars_to_take = min(original_segment_len, len(remaining_text))
                    else:
                        chars_to_take = 1 if remaining_text else 0
                    
                    segment_text = remaining_text[:chars_to_take]
                    remaining_text = remaining_text[chars_to_take:]
                
                if segment_text:  # Only create run if there's text
                    # Clone the XML run element with new text
                    import copy
                    from docx.oxml.ns import qn
                    
                    cloned_run_element = copy.deepcopy(run_element)
                    
                    # Update the text content in all text nodes
                    for t_elem in cloned_run_element.iter(qn('w:t')):
                        t_elem.text = segment_text
                        break  # Only set text in the first text element
                    
                    # Clear text from remaining text elements
                    text_elements = list(cloned_run_element.iter(qn('w:t')))
                    for t_elem in text_elements[1:]:
                        t_elem.text = ""
                    
                    replacement_runs.append(cloned_run_element)
            
            # Handle remaining text using last run's formatting
            while remaining_text and affected_xml_runs:
                last_run_element = affected_xml_runs[-1][0]
                
                import copy
                from docx.oxml.ns import qn
                
                cloned_run_element = copy.deepcopy(last_run_element)
                
                # Update the text content
                for t_elem in cloned_run_element.iter(qn('w:t')):
                    t_elem.text = remaining_text
                    break
                
                # Clear text from remaining text elements
                text_elements = list(cloned_run_element.iter(qn('w:t')))
                for t_elem in text_elements[1:]:
                    t_elem.text = ""
                
                replacement_runs.append(cloned_run_element)
                remaining_text = ""
            
            # Replace the matched XML runs with replacement runs
            if replacement_runs and affected_xml_runs:
                # Insert replacement runs before the first run to replace
                insert_point = affected_xml_runs[0][0]
                for replacement_run in replacement_runs:
                    insert_point.addprevious(replacement_run)
                
                # Remove the original matched runs
                for run_element, _, _, _ in affected_xml_runs:
                    run_element.getparent().remove(run_element)
            
            count += 1
            _dbg("XML path: one replacement applied")
    
    # Process all paragraphs in document body
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process all paragraphs in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    _dbg(f"Total replacements: {count}")
    return count