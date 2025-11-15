"""
Core comment extraction functionality for Word documents.

This module provides low-level functions to extract and process comments
from Word documents using the python-docx library.
"""
import datetime
from typing import Dict, List, Optional, Any
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph


def extract_all_comments(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Extract all comments from a Word document.
    
    Args:
        doc: The Document object to extract comments from
        
    Returns:
        List of dictionaries containing comment information
    """
    comments = []
    
    # Access the document's comment part if it exists
    try:
        # Get the document part
        document_part = doc.part
        
        # Find comments part through relationships
        comments_part = None
        for rel_id, rel in document_part.rels.items():
            if 'comments' in rel.reltype and 'comments' == rel.reltype.split('/')[-1]:
                comments_part = rel.target_part
                break
        
        if comments_part:
            # Extract comments from the comments part using proper xpath syntax
            comment_elements = comments_part.element.xpath('.//w:comment')
            
            for idx, comment_element in enumerate(comment_elements):
                comment_data = extract_comment_data(comment_element, idx)
                if comment_data:
                    comments.append(comment_data)
        
        # If no comments found, try alternative approach
        if not comments:
            # Fallback: scan paragraphs for comment references
            comments = extract_comments_from_paragraphs(doc)
    
    except Exception as e:
        # If direct access fails, try alternative approach
        comments = extract_comments_from_paragraphs(doc)
    
    return comments


def extract_comments_from_paragraphs(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Extract comments by scanning paragraphs for comment references.
    
    Args:
        doc: The Document object
        
    Returns:
        List of comment dictionaries
    """
    comments = []
    comment_id = 1
    
    # Check all paragraphs in the document
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_comments = find_paragraph_comments(paragraph, para_idx, comment_id)
        comments.extend(para_comments)
        comment_id += len(para_comments)
    
    # Check paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    para_comments = find_paragraph_comments(paragraph, para_idx, comment_id, in_table=True)
                    comments.extend(para_comments)
                    comment_id += len(para_comments)
    
    return comments


def extract_comment_data(comment_element, index: int) -> Optional[Dict[str, Any]]:
    """
    Extract data from a comment XML element.
    
    Args:
        comment_element: The XML comment element
        index: Index for generating a unique ID
        
    Returns:
        Dictionary with comment data or None
    """
    try:
        # Extract comment attributes
        comment_id = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(index))
        author = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
        initials = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials', '')
        date_str = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
        
        # Parse date if available
        date = None
        if date_str:
            try:
                date = datetime.datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                date = date.isoformat()
            except:
                date = date_str
        
        # Extract comment text
        text_elements = comment_element.xpath('.//w:t')
        text = ''.join(elem.text or '' for elem in text_elements)
        
        return {
            'id': f'comment_{index + 1}',
            'comment_id': comment_id,
            'author': author,
            'initials': initials,
            'date': date,
            'text': text.strip(),
            'paragraph_index': None,  # Will be filled if we can determine it
            'in_table': False,
            'reference_text': ''
        }
    
    except Exception as e:
        return None


def find_paragraph_comments(paragraph: Paragraph, para_index: int, 
                           start_id: int, in_table: bool = False) -> List[Dict[str, Any]]:
    """
    Find comments associated with a specific paragraph.
    
    Args:
        paragraph: The paragraph to check
        para_index: The index of the paragraph
        start_id: Starting ID for comments
        in_table: Whether the paragraph is in a table
        
    Returns:
        List of comment dictionaries
    """
    comments = []
    
    try:
        # Access the paragraph's XML element
        para_xml = paragraph._element
        
        # Look for comment range markers (simplified approach)
        # This is a basic implementation - the full version would need more sophisticated XML parsing
        xml_text = str(para_xml)
        
        # Simple check for comment references in the XML
        if 'commentRangeStart' in xml_text or 'commentReference' in xml_text:
            # Create a placeholder comment entry
            comment_info = {
                'id': f'comment_{start_id}',
                'comment_id': f'{start_id}',
                'author': 'Unknown',
                'initials': '',
                'date': None,
                'text': 'Comment detected but content not accessible',
                'paragraph_index': para_index,
                'in_table': in_table,
                'reference_text': paragraph.text[:50] + '...' if len(paragraph.text) > 50 else paragraph.text
            }
            comments.append(comment_info)
    
    except Exception:
        # If we can't access the XML, skip this paragraph
        pass
    
    return comments


def filter_comments_by_author(comments: List[Dict[str, Any]], author: str) -> List[Dict[str, Any]]:
    """
    Filter comments by author name.
    
    Args:
        comments: List of comment dictionaries
        author: Author name to filter by (case-insensitive)
        
    Returns:
        Filtered list of comments
    """
    author_lower = author.lower()
    return [c for c in comments if c.get('author', '').lower() == author_lower]


def get_comments_for_paragraph(comments: List[Dict[str, Any]], paragraph_index: int) -> List[Dict[str, Any]]:
    """
    Get all comments for a specific paragraph.
    
    Args:
        comments: List of all comments
        paragraph_index: Index of the paragraph
        
    Returns:
        Comments for the specified paragraph
    """
    return [c for c in comments if c.get('paragraph_index') == paragraph_index]


def add_reply_to_comment(filename: str, comment_id: str, reply_text: str, author: str = "", initials: str = "") -> bool:
    """
    Add a reply to an existing comment in a Word document using COM interface.
    This creates a proper reply that appears as a separate comment in the thread.
    
    Args:
        filename: Path to the Word document
        comment_id: The ID of the comment to reply to (can be numeric string or comment_id from comment data)
        reply_text: Text content for the reply
        author: Author name for the reply (optional)
        initials: Author initials for the reply (optional)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        import win32com.client
        import os
        
        # Convert to absolute path
        abs_filename = os.path.abspath(filename)
        
        if not os.path.exists(abs_filename):
            return False
        
        # FIRST: Load comments using python-docx BEFORE opening via COM API
        # This gives us the stable mapping that doesn't change when we add replies
        target_comment_data = None
        comment_id_str = str(comment_id).strip()
        
        try:
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(abs_filename)
            all_docx_comments = extract_all_comments(docx_doc)
            
            # Filter out replies - we only want main comments
            # Replies can be identified by checking if they have the same text as a known reply pattern
            # But we don't know reply patterns yet, so we'll filter them during matching
            # For now, we'll use all comments but prioritize main comments during matching
            
            # Try to filter obvious replies: comments that have short generic text
            # that matches common reply patterns (like "Отработано")
            # But we need to be careful - some main comments might also be short
            docx_comments = []
            reply_text_lower_check = reply_text.lower().strip() if reply_text else ""
            
            for comm in all_docx_comments:
                comm_text_lower = comm.get('text', '').lower().strip()
                # Skip if this looks like a reply we're trying to add
                # (but only if we're sure it's not the target comment)
                if reply_text_lower_check and comm_text_lower == reply_text_lower_check:
                    # This might be a reply, but we'll include it for now
                    # and filter it out during matching if it doesn't match our target
                    pass
                docx_comments.append(comm)
            
            # Find the comment in our list that matches the comment_id
            # First try exact match by timestamp (if comment_id is a timestamp), then by ID
            reply_text_lower_check = reply_text.lower().strip() if reply_text else ""
            
            # Check if comment_id_str looks like a timestamp (ISO format)
            is_timestamp = False
            try:
                # Try to parse as ISO datetime
                datetime.datetime.fromisoformat(comment_id_str.replace('Z', '+00:00'))
                is_timestamp = True
            except (ValueError, AttributeError):
                pass
            
            # First pass: find by timestamp (if it's a timestamp) or by ID
            candidate_by_id = None
            for comm in docx_comments:
                if is_timestamp:
                    # Match by timestamp
                    comm_date = comm.get('date')
                    if comm_date:
                        # Normalize timestamps for comparison
                        comm_date_str = str(comm_date)
                        comment_id_normalized = comment_id_str.replace('Z', '+00:00')
                        comm_date_normalized = comm_date_str.replace('Z', '+00:00')
                        # Compare timestamps (allowing for slight variations in format)
                        try:
                            comm_dt = datetime.datetime.fromisoformat(comm_date_normalized)
                            target_dt = datetime.datetime.fromisoformat(comment_id_normalized)
                            # Match if timestamps are very close (within 1 second)
                            if abs((comm_dt - target_dt).total_seconds()) < 1:
                                candidate_by_id = comm
                                break
                        except:
                            # Fallback to string comparison
                            if comm_date_str == comment_id_str or comm_date_normalized == comment_id_normalized:
                                candidate_by_id = comm
                                break
                else:
                    # Match by ID
                    comm_matches_id = (comm.get('id') == comment_id_str or
                                       str(comm.get('comment_id')) == comment_id_str)
                    if comm_matches_id:
                        candidate_by_id = comm
                        break
            
            # Check if candidate is a reply and find main comment if needed
            if candidate_by_id:
                comm_text_lower = candidate_by_id.get('text', '').lower().strip()
                
                # Check if this looks like a reply (contains reply text pattern)
                # Replies have the reply text we're trying to add
                # Also check for debug patterns that indicate this is a reply
                is_reply_candidate = False
                if reply_text_lower_check:
                    # Check if text starts with reply text or contains it
                    if comm_text_lower.startswith(reply_text_lower_check) or reply_text_lower_check in comm_text_lower:
                        is_reply_candidate = True
                    # Also check for debug patterns
                    import re
                    if re.search(r'\[к комментарию \d+: comment_\d+\]', comm_text_lower) or \
                       re.search(r'\[COM index=\d+', comm_text_lower):
                        is_reply_candidate = True
                
                # If this is a reply, we need to find the main comment
                if is_reply_candidate:
                    # This is a reply - we need to find the main comment
                    # Strategy: filter out all replies and find main comment by index
                    main_comments_filtered = []
                    for comm in docx_comments:
                        comm_text_check = comm.get('text', '').lower().strip()
                        # Skip replies (comments that contain reply text)
                        if reply_text_lower_check and (comm_text_check.startswith(reply_text_lower_check) or reply_text_lower_check in comm_text_check):
                            continue  # Skip replies
                        main_comments_filtered.append(comm)
                    
                    # Try to find main comment by index from comment_id
                    # The key: if comment_5 is a reply at position 5, the main comment 
                    # should be at position (5 - number_of_replies_before_position_5) in filtered list
                    # But if comment_id is a timestamp, we need a different approach
                    if is_timestamp:
                        # If we found a reply by timestamp, we need to find the main comment
                        # Extract the original comment ID from the reply text if possible
                        reply_text_with_debug = candidate_by_id.get('text', '')
                        # Try to extract comment ID from debug pattern like "[к комментарию X: comment_Y]"
                        import re
                        match = re.search(r'\[к комментарию \d+: (comment_\d+)\]', reply_text_with_debug)
                        if match:
                            original_comment_id = match.group(1)
                            # Find the main comment by this ID
                            for comm in main_comments_filtered:
                                if comm.get('id') == original_comment_id or str(comm.get('comment_id')) == original_comment_id.replace('comment_', ''):
                                    target_comment_data = comm
                                    break
                    elif comment_id_str.startswith('comment_'):
                        try:
                            index_part = int(comment_id_str.replace('comment_', ''))
                            
                            # Count how many replies are before this index in the original list
                            replies_before = 0
                            for i in range(index_part - 1):  # Check comments before this index (0-based)
                                if i < len(docx_comments):
                                    comm_check = docx_comments[i]
                                    comm_text_check = comm_check.get('text', '').lower().strip()
                                    # Check if this is a reply (including debug patterns)
                                    is_reply_check = False
                                    if reply_text_lower_check and (comm_text_check.startswith(reply_text_lower_check) or reply_text_lower_check in comm_text_check):
                                        is_reply_check = True
                                    if re.search(r'\[к комментарию \d+: comment_\d+\]', comm_text_check) or \
                                       re.search(r'\[COM index=\d+', comm_text_check):
                                        is_reply_check = True
                                    if is_reply_check:
                                        replies_before += 1
                            
                            # The main comment should be at position (index_part - 1 - replies_before) in filtered list
                            main_comment_index = index_part - 1 - replies_before
                            if 0 <= main_comment_index < len(main_comments_filtered):
                                target_comment_data = main_comments_filtered[main_comment_index]
                        except (ValueError, TypeError, IndexError) as e:
                            pass
                    
                    # If still not found, try to find by matching characteristics
                    # We know the reply has author "Тестовый пользователь" and contains reply text
                    # The main comment should have a different author
                    if not target_comment_data:
                        reply_author_lower = candidate_by_id.get('author', '').lower()
                        # Try to find main comment by author difference
                        for comm in main_comments_filtered:
                            comm_author_check = comm.get('author', '').lower()
                            # Main comment should have different author than reply
                            if comm_author_check != reply_author_lower:
                                # This could be the main comment
                                if not target_comment_data:
                                    target_comment_data = comm
                                break
                    
                    # If still not found, we can't proceed reliably
                    if not target_comment_data:
                        return False
                else:
                    # Not a reply, use it
                    target_comment_data = candidate_by_id
            
            # If not found by exact match, try to extract index from "comment_N" format
            # But count only MAIN comments (excluding replies) for index matching
            if target_comment_data is None:
                try:
                    if comment_id_str.startswith('comment_'):
                        index_part = int(comment_id_str.replace('comment_', ''))
                        # Filter main comments for index matching
                        main_comments_for_index = []
                        for comm in docx_comments:
                            comm_text_lower = comm.get('text', '').lower().strip()
                            # Skip replies (comments with reply text and different author pattern)
                            if reply_text_lower_check and comm_text_lower == reply_text_lower_check:
                                # This might be a reply - skip it for index matching
                                continue
                            main_comments_for_index.append(comm)
                        
                        if 1 <= index_part <= len(main_comments_for_index):
                            target_comment_data = main_comments_for_index[index_part - 1]  # Convert to 0-based
                    elif comment_id_str.isdigit():
                        index_part = int(comment_id_str)
                        if index_part == 0:
                            index_part = 1
                        # Same filtering for numeric IDs
                        main_comments_for_index = []
                        for comm in docx_comments:
                            comm_text_lower = comm.get('text', '').lower().strip()
                            if reply_text_lower_check and comm_text_lower == reply_text_lower_check:
                                continue
                            main_comments_for_index.append(comm)
                        
                        if 1 <= index_part <= len(main_comments_for_index):
                            target_comment_data = main_comments_for_index[index_part - 1]  # Convert to 0-based
                except (ValueError, TypeError, IndexError):
                    pass
                    
        except Exception as e:
            return False
        
        if target_comment_data is None:
            return False
        
        # Get unique characteristics of the target comment
        target_author = target_comment_data.get('author', '').lower()
        target_text = target_comment_data.get('text', '').strip()
        target_comment_xml_id = str(target_comment_data.get('comment_id', ''))
        target_timestamp = target_comment_data.get('date', None)
        target_index_in_list = None
        
        # Find the index of this comment in the FILTERED list (main comments only)
        # This is critical - we need to count only main comments, not replies
        try:
            # Create filtered list (same filtering as when finding target_comment_data)
            main_comments_for_index_calc = []
            reply_text_lower_check = reply_text.lower().strip() if reply_text else ""
            for comm in docx_comments:
                comm_text_lower = comm.get('text', '').lower().strip()
                if reply_text_lower_check and comm_text_lower == reply_text_lower_check:
                    # Skip replies
                    continue
                main_comments_for_index_calc.append(comm)
            
            # Find index in filtered list
            for idx, comm in enumerate(main_comments_for_index_calc):
                if comm == target_comment_data:
                    target_index_in_list = idx
                    break
        except:
            pass
        
        # Normalize target text for comparison (remove extra whitespace, lowercase)
        target_text_normalized = ' '.join(target_text.lower().split())
        
        # NOW open Word via COM API
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Don't show Word window
        
        try:
            # Open the document
            doc = word.Documents.Open(abs_filename)
            
            # Get comments collection
            comments = doc.Comments
            
            if comments.Count == 0:
                doc.Close(SaveChanges=False)
                word.Quit()
                return False
            
            # Find the comment in COM API by matching unique characteristics
            # We iterate through all comments and match by author and text
            # This approach works even after replies are added, as the main comment
            # retains its original author and text characteristics
            
            # First, build a list of MAIN comments only (excluding replies) from COM API
            # Replies appear as separate comments in the Comments collection and shift indices
            # We need to identify and exclude them to maintain correct mapping
            main_comments_list = []
            for i in range(1, comments.Count + 1):
                try:
                    com_comment = comments.Item(i)
                    com_text = ""
                    com_author = ""
                    try:
                        com_text = com_comment.Range.Text.strip()
                        com_author = com_comment.Author.lower() if hasattr(com_comment, 'Author') else ''
                    except:
                        pass
                    
                    # Check if this comment is a reply
                    # In Word COM API, replies appear as separate comments in the Comments collection
                    # We can't reliably use Parent property as it may exist for all comments
                    # Instead, we'll identify replies by their characteristics:
                    # - They have the same text as a reply we're trying to add
                    # - They have a different author than the original comment
                    is_reply = False
                    
                    # Note: We'll determine if it's a reply during filtering, not here
                    # This flag is kept for potential future use
                    
                    # Method 2: Heuristic - check if comment looks like a reply
                    # Replies often have short generic text and appear after main comments
                    # But we can't rely solely on this, as main comments can also be short
                    # So we'll use this only as a secondary check
                    
                    # For now, we'll use a different approach:
                    # We'll match by author + text, and if we find multiple matches,
                    # we'll prefer the one that matches the expected position better
                    # But first, let's try to exclude obvious replies
                    
                    # Skip if this looks like a reply we just added (heuristic)
                    # This is a workaround - ideally we'd use proper reply detection
                    looks_like_reply = False
                    if com_text.lower() == reply_text.lower().strip():
                        # This comment has the same text as the reply we're trying to add
                        # It's likely a reply that was added previously
                        looks_like_reply = True
                    
                    # Include all comments for now, but mark potential replies
                    # We'll filter them out during matching if they don't match our target
                    main_comments_list.append({
                        'index': i,
                        'comment': com_comment,
                        'author': com_author,
                        'text': com_text,
                        'is_reply': is_reply,
                        'looks_like_reply': looks_like_reply
                    })
                except:
                    pass
            
            target_comment = None
            best_match_score = 0
            best_match_index = None
            
            # Filter out replies from the list - we only want main comments
            # The key insight: replies appear as separate comments in COM API Comments collection
            # We need to identify and exclude them to maintain correct position mapping
            filtered_main_comments = []
            reply_text_lower = reply_text.lower().strip() if reply_text else ""
            
            
            for candidate in main_comments_list:
                candidate_text_lower = candidate['text'].lower()
                
                # Skip if it's identified as a reply via Parent property
                if candidate.get('is_reply', False):
                    continue
                
                # Skip if it looks like a reply we're trying to add
                # A reply contains the reply text we're adding
                if reply_text_lower:
                    # Check if this comment contains reply text
                    contains_reply_text = reply_text_lower in candidate_text_lower or candidate_text_lower.startswith(reply_text_lower)
                    
                    # Also check for debug patterns that indicate this is a reply
                    import re
                    has_debug_pattern = bool(re.search(r'\[к комментарию \d+: comment_\d+\]', candidate_text_lower) or 
                                            re.search(r'\[COM index=\d+', candidate_text_lower))
                    
                    if contains_reply_text or has_debug_pattern:
                        # This comment contains reply text or debug patterns - it's likely a reply
                        # But check if it matches our target (in case target comment itself has this text)
                        if target_author and candidate['author'] != target_author:
                            # Different author - this is definitely a reply, skip it
                            continue
                        # If author matches, check if text also matches target (without reply text)
                        if target_text_normalized:
                            candidate_text_normalized = ' '.join(candidate_text_lower.split())
                            # Remove reply text and debug patterns for comparison
                            candidate_text_clean = re.sub(r'\[к комментарию \d+: comment_\d+\]', '', candidate_text_normalized)
                            candidate_text_clean = re.sub(r'\[COM index=\d+, target index=\d+, target ID=comment_\d+\]', '', candidate_text_clean)
                            candidate_text_clean = candidate_text_clean.replace(reply_text_lower, '').strip()
                            
                            if candidate_text_clean != target_text_normalized and candidate_text_normalized != target_text_normalized:
                                # Text doesn't match target - likely a reply
                                continue
                
                # Include this comment in filtered list (it's a main comment)
                filtered_main_comments.append(candidate)
            
            
            # FIRST: Try to match by timestamp if available (most reliable)
            # But only if target comment is NOT a reply (replies have same timestamp as when they were created)
            if target_timestamp and not target_comment:
                # Check if target comment is a reply by checking if it contains reply text
                target_is_reply = False
                if reply_text_lower:
                    target_text_lower = target_text.lower()
                    if reply_text_lower in target_text_lower or target_text_lower.startswith(reply_text_lower):
                        # This might be a reply - skip timestamp matching and use author+text instead
                        target_is_reply = True
                
                if not target_is_reply:
                    # Match by author and text (timestamp is already verified in python-docx)
                    # This is more reliable than timestamp alone, especially when multiple replies have same timestamp
                    for candidate in filtered_main_comments:
                        candidate_author = candidate['author']
                        candidate_text = candidate['text']
                        
                        # Match by author and text (timestamp is already verified in python-docx)
                        if target_author and candidate_author == target_author:
                            candidate_text_normalized = ' '.join(candidate_text.lower().split())
                            # Remove reply text patterns for comparison
                            import re
                            candidate_text_clean = re.sub(r'\[к комментарию \d+: comment_\d+\]', '', candidate_text_normalized).strip()
                            candidate_text_clean = re.sub(r'\[COM index=\d+.*?\]', '', candidate_text_clean).strip()
                            
                            if target_text_normalized == candidate_text_normalized or \
                               (len(target_text_normalized) >= 20 and candidate_text_clean.startswith(target_text_normalized[:20])) or \
                               (len(target_text_normalized) >= 20 and candidate_text_normalized.startswith(target_text_normalized[:20])):
                                target_comment = candidate['comment']
                                best_match_index = candidate['index']
                                best_match_score = 200  # Very high score for timestamp+author+text match
                                break
            
            # SECOND: Try to match by position in the filtered list (main comments only)
            # This accounts for replies that shift the indices
            # IMPORTANT: target_index_in_list is the index in filtered_main_comments (python-docx)
            # We need to find the corresponding comment in filtered_main_comments (COM API)
            if target_index_in_list is not None and target_index_in_list < len(filtered_main_comments) and not target_comment:
                candidate = filtered_main_comments[target_index_in_list]
                candidate_author = candidate['author']
                candidate_text_normalized = ' '.join(candidate['text'].lower().split())
                
                # Remove reply text from candidate text for comparison (if it contains reply text)
                if reply_text_lower and reply_text_lower in candidate_text_normalized:
                    # This might be a reply - extract the original text
                    # Try to remove reply text pattern to get original text
                    candidate_text_original = candidate_text_normalized.replace(reply_text_lower, '').strip()
                    # Remove debug info patterns like "[к комментарию X: comment_Y]"
                    import re
                    candidate_text_original = re.sub(r'\[к комментарию \d+: comment_\d+\]', '', candidate_text_original).strip()
                    candidate_text_original = re.sub(r'\[COM index=\d+, target index=\d+, target ID=comment_\d+\]', '', candidate_text_original).strip()
                    if candidate_text_original:
                        candidate_text_normalized = candidate_text_original
                
                # Check if this candidate matches our target
                author_match = not target_author or candidate_author == target_author
                text_match = False
                
                if target_text_normalized and candidate_text_normalized:
                    # Exact match or very close
                    if target_text_normalized == candidate_text_normalized:
                        text_match = True
                    elif len(target_text_normalized) >= 20 and len(candidate_text_normalized) >= 20:
                        # Check if first 20 chars match (for longer texts)
                        if target_text_normalized[:20] == candidate_text_normalized[:20]:
                            text_match = True
                    elif len(target_text_normalized) < 20:
                        # For short texts, require exact match
                        text_match = target_text_normalized == candidate_text_normalized
                
                if author_match and (text_match or not target_text_normalized):
                    # Position-based match found in filtered list!
                    target_comment = candidate['comment']
                    best_match_index = candidate['index']
                    best_match_score = 100  # Very high score for position match
            
            # If position-based match didn't work, try matching by characteristics
            # Use filtered list to avoid matching replies
            if not target_comment:
                for candidate in filtered_main_comments:
                    i = candidate['index']
                    com_comment = candidate['comment']
                    com_author = candidate['author']
                    com_text = candidate['text']
                    com_text_normalized = ' '.join(com_text.lower().split())
                    
                    # Author must match
                    if target_author and com_author != target_author:
                        continue
                    
                    # Calculate match score (higher is better)
                    match_score = 0
                    
                    # Author match (required)
                    if target_author and com_author == target_author:
                        match_score += 10
                    else:
                        continue  # Skip if author doesn't match
                    
                    # Text match is CRITICAL - we need exact or very close match
                    if target_text_normalized and com_text_normalized:
                        # Calculate similarity
                        target_len = len(target_text_normalized)
                        com_len = len(com_text_normalized)
                        
                        # Exact match
                        if target_text_normalized == com_text_normalized:
                            match_score += 50  # Very high score for exact match
                        # Check if one contains the other (for cases where text might be truncated)
                        elif target_text_normalized in com_text_normalized:
                            # Check how much of target text is in COM text
                            overlap_ratio = len(target_text_normalized) / max(com_len, 1)
                            if overlap_ratio > 0.9:  # At least 90% overlap (more strict)
                                match_score += 40
                            elif overlap_ratio > 0.7:  # At least 70% overlap
                                match_score += 25
                            else:
                                # Not enough overlap
                                continue
                        elif com_text_normalized in target_text_normalized:
                            overlap_ratio = len(com_text_normalized) / max(target_len, 1)
                            if overlap_ratio > 0.9:
                                match_score += 40
                            elif overlap_ratio > 0.7:
                                match_score += 25
                            else:
                                continue
                        # Check prefix match (first 50+ chars) - only for longer texts
                        elif target_len >= 50 and com_len >= 50:
                            if target_text_normalized[:50] == com_text_normalized[:50]:
                                match_score += 30
                            elif target_text_normalized[:30] == com_text_normalized[:30]:
                                match_score += 20
                            else:
                                # Prefix doesn't match well enough
                                continue
                        else:
                            # Short texts - require exact match
                            if target_text_normalized == com_text_normalized:
                                match_score += 20
                            else:
                                # Text doesn't match - skip
                                continue
                    elif not target_text_normalized and not com_text_normalized:
                        # Both empty - still match by author
                        match_score += 5
                    else:
                        # One has text, other doesn't - not a good match
                        continue
                    
                    # Bonus: if index is close to expected position
                    if target_index_in_list is not None:
                        expected_com_index = target_index_in_list + 1
                        index_diff = abs(i - expected_com_index)
                        if index_diff == 0:
                            match_score += 10  # Exact position match
                        elif index_diff <= 2:
                            match_score += 5  # Close position
                    
                    # If this is a better match than previous, save it
                    if match_score > best_match_score:
                        best_match_score = match_score
                        target_comment = com_comment
                        best_match_index = i
                        
                        # If we have a very high match score (exact author + exact text), stop searching
                        if match_score >= 60:
                            break
            
            if not target_comment:
                doc.Close(SaveChanges=False)
                word.Quit()
                return False
            
            # Add reply to the comment
            # Replies.Add requires a Range object - use the comment's Range
            reply = target_comment.Replies.Add(Range=target_comment.Range, Text=reply_text)
            
            # Set author if provided
            if author:
                reply.Author = author
            
            # Set initials if provided
            if initials:
                reply.Initial = initials
            
            # Save and close
            doc.Save()
            doc.Close()
            word.Quit()
            
            return True
            
        except Exception as e:
            try:
                doc.Close(SaveChanges=False)
            except:
                pass
            try:
                word.Quit()
            except:
                pass
            return False
            
    except ImportError:
        return False
    except Exception as e:
        return False