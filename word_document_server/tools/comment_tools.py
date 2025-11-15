"""
Comment extraction tools for Word Document Server.

These tools provide high-level interfaces for extracting and analyzing
comments from Word documents through the MCP protocol.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import ensure_docx_extension
from word_document_server.core.comments import (
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph,
    add_reply_to_comment
)


async def get_all_comments(filename: str) -> str:
    """
    Extract all comments from a Word document.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string containing all comments with metadata
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        comments = extract_all_comments(doc)
        
        # Return results
        return json.dumps({
            'success': True,
            'comments': comments,
            'total_comments': len(comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_by_author(filename: str, author: str) -> str:
    """
    Extract comments from a specific author in a Word document.
    
    Args:
        filename: Path to the Word document
        author: Name of the comment author to filter by
        
    Returns:
        JSON string containing filtered comments
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if not author or not author.strip():
        return json.dumps({
            'success': False,
            'error': 'Author name cannot be empty'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter by author
        author_comments = filter_comments_by_author(all_comments, author)
        
        # Return results
        return json.dumps({
            'success': True,
            'author': author,
            'comments': author_comments,
            'total_comments': len(author_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_for_paragraph(filename: str, paragraph_index: int) -> str:
    """
    Extract comments for a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        
    Returns:
        JSON string containing comments for the specified paragraph
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if paragraph_index < 0:
        return json.dumps({
            'success': False,
            'error': 'Paragraph index must be non-negative'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Check if paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                'success': False,
                'error': f'Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.'
            }, indent=2)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter for the specific paragraph
        from word_document_server.core.comments import get_comments_for_paragraph as core_get_comments_for_paragraph
        para_comments = core_get_comments_for_paragraph(all_comments, paragraph_index)
        
        # Get the paragraph text for context
        paragraph_text = doc.paragraphs[paragraph_index].text
        
        # Return results
        return json.dumps({
            'success': True,
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph_text,
            'comments': para_comments,
            'total_comments': len(para_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def reply_to_comment(filename: str, comment_id: str, reply_text: str, author: str = "", initials: str = "") -> str:
    """
    Add a reply to an existing comment in a Word document.
    
    Args:
        filename: Path to the Word document
        comment_id: Identifier of the comment to reply to. Can be:
            - Timestamp (ISO format, e.g., "2025-11-06T21:57:00+00:00") - recommended for reliable identification
            - Comment ID from comment data (e.g., "comment_1", "comment_2")
            - Numeric string (e.g., "0", "1") - less reliable as indices may shift
        reply_text: Text content for the reply
        author: Author name for the reply (optional)
        initials: Author initials for the reply (optional)
        
    Returns:
        JSON string indicating success or failure
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if not comment_id or not str(comment_id).strip():
        return json.dumps({
            'success': False,
            'error': 'Comment ID cannot be empty'
        }, indent=2)
    
    if not reply_text or not reply_text.strip():
        return json.dumps({
            'success': False,
            'error': 'Reply text cannot be empty'
        }, indent=2)
    
    try:
        # Add reply to the comment using COM interface
        # This creates a proper reply that appears as a separate comment in the thread
        success = add_reply_to_comment(filename, comment_id, reply_text, author, initials)
        
        if success:
            return json.dumps({
                'success': True,
                'message': f'Reply added to comment {comment_id}',
                'comment_id': comment_id,
                'reply_text': reply_text,
                'author': author if author else 'Current user'
            }, indent=2)
        else:
            return json.dumps({
                'success': False,
                'error': f'Comment with ID {comment_id} not found or could not be accessed. Make sure pywin32 is installed.'
            }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to add reply to comment: {str(e)}'
        }, indent=2)