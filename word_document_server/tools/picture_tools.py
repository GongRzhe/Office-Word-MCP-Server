"""
Picture tools for Word Document Server.

These tools modify existing pictures/images in Word documents,
including resizing and alignment.
"""
import os
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


async def resize_picture(filename: str, picture_index: int, width: Optional[float] = None, 
                        height: Optional[float] = None, maintain_aspect_ratio: bool = True) -> str:
    """Resize an existing picture in a Word document.
    
    Args:
        filename: Path to the Word document
        picture_index: Index of the picture to resize (0-based)
        width: New width in inches (optional)
        height: New height in inches (optional)
        maintain_aspect_ratio: If True and only width or height is specified, 
                              the other dimension will be scaled proportionally
    
    Returns:
        Status message indicating success or failure
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        picture_index = int(picture_index)
        if width is not None:
            width = float(width)
        if height is not None:
            height = float(height)
    except (ValueError, TypeError):
        return "Invalid parameter: picture_index must be an integer, width and height must be numbers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    if width is None and height is None:
        return "Error: At least one of width or height must be specified"
    
    try:
        doc = Document(filename)
        
        # Find all pictures in the document
        pictures = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check if run contains a picture
                if 'graphic' in run._element.xml:
                    # Get inline shapes (pictures)
                    inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for inline_shape in inline_shapes:
                        pictures.append((run, inline_shape))
        
        # Also check in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'graphic' in run._element.xml:
                                inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                for inline_shape in inline_shapes:
                                    pictures.append((run, inline_shape))
        
        if not pictures:
            return f"No pictures found in {filename}"
        
        if picture_index < 0 or picture_index >= len(pictures):
            return f"Invalid picture index {picture_index}. Document contains {len(pictures)} picture(s) (0-{len(pictures)-1})"
        
        run, inline_shape = pictures[picture_index]
        
        # Get the extent element which contains the size
        extent = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
        if extent is None:
            return f"Could not find size information for picture {picture_index}"
        
        # Get current size in EMUs (English Metric Units)
        # 914400 EMUs = 1 inch
        current_width_emu = int(extent.get('cx'))
        current_height_emu = int(extent.get('cy'))
        
        # Convert current size to inches
        current_width_inches = current_width_emu / 914400
        current_height_inches = current_height_emu / 914400
        
        # Calculate new size
        if maintain_aspect_ratio:
            if width is not None and height is None:
                # Calculate height based on width
                aspect_ratio = current_height_inches / current_width_inches
                new_width_inches = width
                new_height_inches = width * aspect_ratio
            elif height is not None and width is None:
                # Calculate width based on height
                aspect_ratio = current_width_inches / current_height_inches
                new_height_inches = height
                new_width_inches = height * aspect_ratio
            else:
                # Both specified, use as is
                new_width_inches = width
                new_height_inches = height
        else:
            # No aspect ratio maintenance
            new_width_inches = width if width is not None else current_width_inches
            new_height_inches = height if height is not None else current_height_inches
        
        # Convert to EMUs
        new_width_emu = int(new_width_inches * 914400)
        new_height_emu = int(new_height_inches * 914400)
        
        # Update the extent element
        extent.set('cx', str(new_width_emu))
        extent.set('cy', str(new_height_emu))
        
        # Also update the inline shape's extent if it exists
        inline = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if inline is not None:
            # Update the graphic's extents as well
            graphic = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
            if graphic is not None:
                graphic_data = graphic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData')
                if graphic_data is not None:
                    pic = graphic_data.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
                    if pic is not None:
                        sp_pr = pic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
                        if sp_pr is not None:
                            xfrm = sp_pr.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm')
                            if xfrm is not None:
                                ext = xfrm.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ext')
                                if ext is not None:
                                    ext.set('cx', str(new_width_emu))
                                    ext.set('cy', str(new_height_emu))
        
        doc.save(filename)
        
        return f"Picture {picture_index} resized successfully from {current_width_inches:.2f}\"x{current_height_inches:.2f}\" to {new_width_inches:.2f}\"x{new_height_inches:.2f}\""
        
    except Exception as e:
        return f"Failed to resize picture: {str(e)}"


async def align_picture(filename: str, picture_index: int, alignment: str = "center") -> str:
    """Set the alignment of a picture in a Word document.
    
    Args:
        filename: Path to the Word document
        picture_index: Index of the picture to align (0-based)
        alignment: Alignment type - "left", "center", "right", or "justify"
    
    Returns:
        Status message indicating success or failure
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        picture_index = int(picture_index)
    except (ValueError, TypeError):
        return "Invalid parameter: picture_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    # Validate alignment
    alignment = alignment.lower()
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    if alignment not in alignment_map:
        return f"Invalid alignment '{alignment}'. Must be one of: left, center, right, justify"
    
    try:
        doc = Document(filename)
        
        # Find all pictures and their paragraphs
        picture_paragraphs = []
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if 'graphic' in run._element.xml:
                    inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if inline_shapes:
                        picture_paragraphs.append(paragraph)
                        break  # Only count this paragraph once
        
        # Also check in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'graphic' in run._element.xml:
                                inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                if inline_shapes:
                                    picture_paragraphs.append(paragraph)
                                    break
        
        if not picture_paragraphs:
            return f"No pictures found in {filename}"
        
        if picture_index < 0 or picture_index >= len(picture_paragraphs):
            return f"Invalid picture index {picture_index}. Document contains {len(picture_paragraphs)} picture(s) (0-{len(picture_paragraphs)-1})"
        
        target_paragraph = picture_paragraphs[picture_index]
        target_paragraph.alignment = alignment_map[alignment]
        
        doc.save(filename)
        
        return f"Picture {picture_index} aligned to {alignment} successfully"
        
    except Exception as e:
        return f"Failed to align picture: {str(e)}"


async def list_pictures(filename: str) -> str:
    """List all pictures in a Word document with their current properties.
    
    Args:
        filename: Path to the Word document
    
    Returns:
        Information about all pictures in the document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        
        pictures_info = []
        picture_count = 0
        
        # Check paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if 'graphic' in run._element.xml:
                    inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for inline_shape in inline_shapes:
                        extent = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                        if extent is not None:
                            width_emu = int(extent.get('cx'))
                            height_emu = int(extent.get('cy'))
                            width_inches = width_emu / 914400
                            height_inches = height_emu / 914400
                            
                            alignment = paragraph.alignment
                            alignment_str = "undefined"
                            if alignment == WD_ALIGN_PARAGRAPH.LEFT:
                                alignment_str = "left"
                            elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                alignment_str = "center"
                            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                alignment_str = "right"
                            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                alignment_str = "justify"
                            
                            pictures_info.append({
                                "index": picture_count,
                                "location": f"paragraph {para_idx}",
                                "width_inches": round(width_inches, 2),
                                "height_inches": round(height_inches, 2),
                                "alignment": alignment_str
                            })
                            picture_count += 1
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run in paragraph.runs:
                            if 'graphic' in run._element.xml:
                                inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                for inline_shape in inline_shapes:
                                    extent = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                    if extent is not None:
                                        width_emu = int(extent.get('cx'))
                                        height_emu = int(extent.get('cy'))
                                        width_inches = width_emu / 914400
                                        height_inches = height_emu / 914400
                                        
                                        alignment = paragraph.alignment
                                        alignment_str = "undefined"
                                        if alignment == WD_ALIGN_PARAGRAPH.LEFT:
                                            alignment_str = "left"
                                        elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                            alignment_str = "center"
                                        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                            alignment_str = "right"
                                        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                            alignment_str = "justify"
                                        
                                        pictures_info.append({
                                            "index": picture_count,
                                            "location": f"table {table_idx}, row {row_idx}, cell {cell_idx}",
                                            "width_inches": round(width_inches, 2),
                                            "height_inches": round(height_inches, 2),
                                            "alignment": alignment_str
                                        })
                                        picture_count += 1
        
        if not pictures_info:
            return f"No pictures found in {filename}"
        
        # Format the result
        result = f"Found {len(pictures_info)} picture(s) in {filename}:\n\n"
        for pic in pictures_info:
            result += f"Picture {pic['index']}:\n"
            result += f"  Location: {pic['location']}\n"
            result += f"  Size: {pic['width_inches']}\" x {pic['height_inches']}\"\n"
            result += f"  Alignment: {pic['alignment']}\n\n"
        
        return result.strip()
        
    except Exception as e:
        return f"Failed to list pictures: {str(e)}"


async def align_pictures_batch(filename: str, picture_indices: list, alignment: str = "center") -> str:
    """Align multiple pictures at once by specifying their indices.
    
    Args:
        filename: Path to the Word document
        picture_indices: List of picture indices to align (0-based)
        alignment: Alignment type - "left", "center", "right", or "justify"
    
    Returns:
        Status message with results for each picture
    """
    filename = ensure_docx_extension(filename)
    
    # Validate alignment
    alignment = alignment.lower()
    valid_alignments = ["left", "center", "right", "justify"]
    if alignment not in valid_alignments:
        return f"Invalid alignment '{alignment}'. Must be one of: {', '.join(valid_alignments)}"
    
    # Validate picture_indices is a list
    if not isinstance(picture_indices, list):
        return "Invalid parameter: picture_indices must be a list of integers"
    
    if not picture_indices:
        return "Error: picture_indices list is empty"
    
    # Convert all indices to integers
    try:
        picture_indices = [int(idx) for idx in picture_indices]
    except (ValueError, TypeError):
        return "Invalid parameter: all picture indices must be integers"
    
    # Process each picture
    results = []
    success_count = 0
    failed_count = 0
    
    for idx in picture_indices:
        result = await align_picture(filename, idx, alignment)
        if "successfully" in result.lower():
            success_count += 1
            results.append(f"✓ Picture {idx}: {result}")
        else:
            failed_count += 1
            results.append(f"✗ Picture {idx}: {result}")
    
    # Build summary
    summary = f"Batch alignment completed: {success_count} succeeded, {failed_count} failed\n\n"
    summary += "\n".join(results)
    
    return summary


async def align_all_pictures(filename: str, alignment: str = "center") -> str:
    """Align all pictures in the document to the same alignment.
    
    Args:
        filename: Path to the Word document
        alignment: Alignment type - "left", "center", "right", or "justify"
    
    Returns:
        Status message with results for all pictures
    """
    filename = ensure_docx_extension(filename)
    
    # First, list all pictures to get their indices
    list_result = await list_pictures(filename)
    
    if "No pictures found" in list_result:
        return list_result
    
    # Parse the number of pictures from the result
    import re
    match = re.search(r"Found (\d+) picture\(s\)", list_result)
    if not match:
        return "Failed to determine number of pictures in document"
    
    picture_count = int(match.group(1))
    
    # Create a list of all indices
    all_indices = list(range(picture_count))
    
    # Use batch alignment
    return await align_pictures_batch(filename, all_indices, alignment)


async def resize_pictures_batch(filename: str, picture_indices: list, width: Optional[float] = None,
                                height: Optional[float] = None, maintain_aspect_ratio: bool = True) -> str:
    """Resize multiple pictures at once by specifying their indices.
    
    Args:
        filename: Path to the Word document
        picture_indices: List of picture indices to resize (0-based)
        width: New width in inches (optional)
        height: New height in inches (optional)
        maintain_aspect_ratio: If True and only width or height is specified,
                              the other dimension will be scaled proportionally
    
    Returns:
        Status message with results for each picture
    """
    filename = ensure_docx_extension(filename)
    
    # Validate picture_indices is a list
    if not isinstance(picture_indices, list):
        return "Invalid parameter: picture_indices must be a list of integers"
    
    if not picture_indices:
        return "Error: picture_indices list is empty"
    
    # Convert all indices to integers
    try:
        picture_indices = [int(idx) for idx in picture_indices]
    except (ValueError, TypeError):
        return "Invalid parameter: all picture indices must be integers"
    
    if width is None and height is None:
        return "Error: At least one of width or height must be specified"
    
    # Process each picture
    results = []
    success_count = 0
    failed_count = 0
    
    for idx in picture_indices:
        result = await resize_picture(filename, idx, width, height, maintain_aspect_ratio)
        if "successfully" in result.lower():
            success_count += 1
            results.append(f"✓ Picture {idx}: {result}")
        else:
            failed_count += 1
            results.append(f"✗ Picture {idx}: {result}")
    
    # Build summary
    summary = f"Batch resize completed: {success_count} succeeded, {failed_count} failed\n\n"
    summary += "\n".join(results)
    
    return summary


async def resize_all_pictures(filename: str, width: Optional[float] = None,
                              height: Optional[float] = None, maintain_aspect_ratio: bool = True) -> str:
    """Resize all pictures in the document to the same dimensions.
    
    Args:
        filename: Path to the Word document
        width: New width in inches (optional)
        height: New height in inches (optional)
        maintain_aspect_ratio: If True and only width or height is specified,
                              the other dimension will be scaled proportionally
    
    Returns:
        Status message with results for all pictures
    """
    filename = ensure_docx_extension(filename)
    
    # First, list all pictures to get their indices
    list_result = await list_pictures(filename)
    
    if "No pictures found" in list_result:
        return list_result
    
    # Parse the number of pictures from the result
    import re
    match = re.search(r"Found (\d+) picture\(s\)", list_result)
    if not match:
        return "Failed to determine number of pictures in document"
    
    picture_count = int(match.group(1))
    
    # Create a list of all indices
    all_indices = list(range(picture_count))
    
    # Use batch resize
    return await resize_pictures_batch(filename, all_indices, width, height, maintain_aspect_ratio)


async def process_pictures_by_size(filename: str, min_width: Optional[float] = None,
                                   max_width: Optional[float] = None, min_height: Optional[float] = None,
                                   max_height: Optional[float] = None, alignment: Optional[str] = None,
                                   resize_width: Optional[float] = None, resize_height: Optional[float] = None) -> str:
    """Process pictures based on size filters - align or resize pictures that match criteria.
    
    Args:
        filename: Path to the Word document
        min_width: Minimum width in inches (pictures >= this will be processed)
        max_width: Maximum width in inches (pictures <= this will be processed)
        min_height: Minimum height in inches (pictures >= this will be processed)
        max_height: Maximum height in inches (pictures <= this will be processed)
        alignment: Optional alignment to apply ("left", "center", "right", "justify")
        resize_width: Optional new width in inches
        resize_height: Optional new height in inches
    
    Returns:
        Status message with results for filtered pictures
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        
        # Get all pictures with their info
        pictures = []
        picture_count = 0
        
        # Check paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if 'graphic' in run._element.xml:
                    inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for inline_shape in inline_shapes:
                        extent = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                        if extent is not None:
                            width_emu = int(extent.get('cx'))
                            height_emu = int(extent.get('cy'))
                            width_inches = width_emu / 914400
                            height_inches = height_emu / 914400
                            
                            pictures.append({
                                "index": picture_count,
                                "width": width_inches,
                                "height": height_inches
                            })
                            picture_count += 1
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'graphic' in run._element.xml:
                                inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                for inline_shape in inline_shapes:
                                    extent = inline_shape.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                    if extent is not None:
                                        width_emu = int(extent.get('cx'))
                                        height_emu = int(extent.get('cy'))
                                        width_inches = width_emu / 914400
                                        height_inches = height_emu / 914400
                                        
                                        pictures.append({
                                            "index": picture_count,
                                            "width": width_inches,
                                            "height": height_inches
                                        })
                                        picture_count += 1
        
        if not pictures:
            return f"No pictures found in {filename}"
        
        # Filter pictures based on criteria
        filtered_indices = []
        for pic in pictures:
            matches = True
            
            if min_width is not None and pic["width"] < min_width:
                matches = False
            if max_width is not None and pic["width"] > max_width:
                matches = False
            if min_height is not None and pic["height"] < min_height:
                matches = False
            if max_height is not None and pic["height"] > max_height:
                matches = False
            
            if matches:
                filtered_indices.append(pic["index"])
        
        if not filtered_indices:
            criteria = []
            if min_width is not None:
                criteria.append(f"width >= {min_width}\"")
            if max_width is not None:
                criteria.append(f"width <= {max_width}\"")
            if min_height is not None:
                criteria.append(f"height >= {min_height}\"")
            if max_height is not None:
                criteria.append(f"height <= {max_height}\"")
            
            return f"No pictures match the criteria: {', '.join(criteria)}"
        
        # Apply operations
        results = []
        
        if alignment:
            align_result = await align_pictures_batch(filename, filtered_indices, alignment)
            results.append(f"Alignment results:\n{align_result}")
        
        if resize_width is not None or resize_height is not None:
            resize_result = await resize_pictures_batch(filename, filtered_indices, resize_width, resize_height, True)
            results.append(f"Resize results:\n{resize_result}")
        
        if not alignment and resize_width is None and resize_height is None:
            return f"Found {len(filtered_indices)} matching picture(s): {filtered_indices}. Please specify alignment or resize parameters to process them."
        
        summary = f"Processed {len(filtered_indices)} picture(s) matching criteria: indices {filtered_indices}\n\n"
        summary += "\n\n".join(results)
        
        return summary
        
    except Exception as e:
        return f"Failed to process pictures by size: {str(e)}"

